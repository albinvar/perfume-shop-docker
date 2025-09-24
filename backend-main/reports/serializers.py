from rest_framework import serializers
from .models import Report
from purchases.models import Purchase, PurchaseItem
from Sales.models import Sale, SaleItem
from django.db.models import Sum, Count
from datetime import datetime
import pandas as pd
from io import BytesIO
from django.http import HttpResponse
import json
from django.db.models.functions import TruncDate
from django.db.models import F, ExpressionWrapper, DecimalField
import os
from django.conf import settings
from django.db import models
import logging
import cloudinary.uploader

logger = logging.getLogger(__name__)

class ReportSerializer(serializers.ModelSerializer):
    file_url = serializers.URLField(read_only=True)
    store_name = serializers.SerializerMethodField()
    created_by_username = serializers.SerializerMethodField()
    
    class Meta:
        model = Report
        fields = ['id', 'report_type', 'format', 'start_date', 'end_date', 
                 'store', 'store_name', 'created_by', 'created_by_username',
                 'created_at', 'file_url']
        read_only_fields = ['created_by', 'created_at', 'file_url']
    
    def get_store_name(self, obj):
        return obj.store.name if obj.store else None
    
    def get_created_by_username(self, obj):
        return obj.created_by.username

class ReportGeneratorSerializer(serializers.Serializer):
    report_type = serializers.ChoiceField(choices=Report.REPORT_TYPES)
    format = serializers.ChoiceField(choices=Report.FORMATS)
    start_date = serializers.DateField()
    end_date = serializers.DateField()
    store = serializers.IntegerField(required=False, allow_null=True)

    def validate(self, data):
        if data['start_date'] > data['end_date']:
            raise serializers.ValidationError("End date must be after start date")
        return data

    def generate_report(self, user):
        validated_data = self.validated_data
        report_type = validated_data['report_type']
        file_format = validated_data['format']
        start_date = validated_data['start_date']
        end_date = validated_data['end_date']
        store_id = validated_data.get('store')

        try:
            has_data = self._check_data_exists(report_type, start_date, end_date, store_id, user)
            if not has_data:
                raise serializers.ValidationError("No data available for the selected criteria")
            
            file_content = self._generate_file_content(report_type, file_format, start_date, end_date, store_id, user)
            
            if file_content is None:
                raise serializers.ValidationError("No data available for the selected criteria")
            
            # Upload to Cloudinary
            filename = f"{report_type.lower()}_report_{start_date}_{end_date}.{file_format.lower()}"
            result = cloudinary.uploader.upload(
                file_content,
                resource_type='raw',
                folder='reports/',
                public_id=filename.split('.')[0],
                format=file_format.lower()
            )
            
            # Create report with the Cloudinary URL
            report = Report.objects.create(
                report_type=report_type,
                format=file_format,
                start_date=start_date,
                end_date=end_date,
                store_id=store_id,
                created_by=user,
                file_url=result['secure_url']
            )
            
            return report
        except Exception as e:
            logger.error(f"Error generating report: {str(e)}", exc_info=True)
            raise serializers.ValidationError(f"Error generating report: {str(e)}")

    def _check_data_exists(self, report_type, start_date, end_date, store_id, user):
        try:
            if report_type == 'PURCHASE':
                qs = Purchase.objects.filter(
                    date__range=[start_date, end_date],
                    is_return=False
                )
            elif report_type == 'PURCHASE_RETURN':
                qs = Purchase.objects.filter(
                    date__range=[start_date, end_date],
                    is_return=True
                )
            elif report_type == 'SALE':
                qs = Sale.objects.filter(
                    date__range=[start_date, end_date],
                    is_return=False
                )
            elif report_type == 'SALE_RETURN':
                qs = Sale.objects.filter(
                    date__range=[start_date, end_date],
                    is_return=True
                )
            else:
                return False
            
            if store_id:
                qs = qs.filter(store_id=store_id)
            elif user.role == 'STAFF' and user.store:
                qs = qs.filter(store=user.store)
            
            return qs.exists()
        except Exception as e:
            logger.error(f"Error checking data existence: {str(e)}")
            return False

    def _generate_file_content(self, report_type, file_format, start_date, end_date, store_id, user):
        try:
            if report_type == 'PURCHASE':
                df = self._get_purchase_data(start_date, end_date, store_id, user, is_return=False)
            elif report_type == 'PURCHASE_RETURN':
                df = self._get_purchase_data(start_date, end_date, store_id, user, is_return=True)
            elif report_type == 'SALE':
                df = self._get_sale_data(start_date, end_date, store_id, user, is_return=False)
            elif report_type == 'SALE_RETURN':
                df = self._get_sale_data(start_date, end_date, store_id, user, is_return=True)
            else:
                raise serializers.ValidationError("Invalid report type")
            
            if df.empty:
                return None
            return self._export_to_format(df, file_format, f"{report_type.replace('_', ' ')} Report")
        except Exception as e:
            logger.error(f"Report generation failed: {str(e)}", exc_info=True)
            raise serializers.ValidationError(f"Report generation failed: {str(e)}")

    def _get_purchase_data(self, start_date, end_date, store_id, user, is_return=False):
        purchase_qs = Purchase.objects.filter(
            date__range=[start_date, end_date],
            is_return=is_return
        ).select_related('store', 'supplier', 'created_by')
        
        if store_id:
            purchase_qs = purchase_qs.filter(store_id=store_id)
        elif user.role == 'STAFF' and user.store:
            purchase_qs = purchase_qs.filter(store=user.store)
        
        purchase_items = PurchaseItem.objects.filter(
            purchase__in=purchase_qs
        ).select_related('product', 'purchase')
        
        data = []
        for idx, item in enumerate(purchase_items, 1):
            purchase = item.purchase
            purchase_date = purchase.date.strftime('%Y-%m-%d') if purchase.date else ''
            
            row_data = {
                'SI No': idx,
                'Date': purchase_date,
                'Type': 'Purchase Return' if is_return else 'Purchase',
                'Invoice No': purchase.invoice_no,
                'Store': purchase.store.name if purchase.store else '',
                'Supplier': purchase.supplier.name if purchase.supplier else '',
                'Created By': purchase.created_by.username if purchase.created_by else '',
                'Product': item.product.name if item.product else '',
                'Product Code': item.product.product_code if item.product else '',
                'Quantity': float(item.quantity) if item.quantity else 0.0,
                'Rate': float(item.rate) if item.rate else 0.0,
                'Amount': float(item.amount) if item.amount else 0.0,
            }
            
            if is_return:
                row_data['Original Invoice'] = purchase.return_reference.invoice_no if purchase.return_reference else ''
            
            data.append(row_data)
        
        columns = [
            'SI No', 'Date', 'Type', 'Invoice No', 'Store', 'Supplier', 'Created By',
            'Product', 'Product Code', 'Quantity', 'Rate', 'Amount'
        ]
        
        if is_return:
            columns.insert(4, 'Original Invoice')
        
        df = pd.DataFrame(data, columns=columns)
        return df

    def _get_sale_data(self, start_date, end_date, store_id, user, is_return=False):
        sale_qs = Sale.objects.filter(
            date__range=[start_date, end_date],
            is_return=is_return
        ).select_related('store', 'created_by', 'customer')
        
        if store_id:
            sale_qs = sale_qs.filter(store_id=store_id)
        elif user.role == 'STAFF' and user.store:
            sale_qs = sale_qs.filter(store=user.store)
        
        sale_items = SaleItem.objects.filter(
            sale__in=sale_qs
        ).select_related('product', 'sale')
        
        data = []
        for idx, item in enumerate(sale_items, 1):
            sale = item.sale
            sale_date = sale.date.strftime('%Y-%m-%d') if sale.date else ''
            
            row_data = {
                'SI No': idx,
                'Date': sale_date,
                'Type': 'Sale Return' if is_return else 'Sale',
                'Invoice No': sale.invoice_no,
                'Store': sale.store.name if sale.store else '',
                'Customer': sale.customer.name if sale.customer else 'Walk-in',
                'Created By': sale.created_by.username if sale.created_by else '',
                'Product': item.product.name if item.product else '',
                'Product Code': item.product.product_code if item.product else '',
                'Quantity': float(item.quantity) if item.quantity else 0.0,
                'Rate': float(item.rate) if item.rate else 0.0,
                'Tax Type': item.tax_type if item.tax_type else '',
                'Tax 1 %': float(item.tax1_rate) if item.tax1_rate else 0.0,
                'Tax 2 %': float(item.tax2_rate) if item.tax2_rate else 0.0,
                'Tax Amount': float(sale.tax_amount) if sale.tax_amount else 0.0,
                'Discount Amount': float(sale.discount_amount) if sale.discount_amount else 0.0,
                'Amount': float(item.amount) if item.amount else 0.0,
                'Final Amount': float(sale.final_amount) if sale.final_amount else 0.0,
            }
            
            if is_return:
                row_data['Original Invoice'] = sale.original_sale.invoice_no if sale.original_sale else ''
            
            data.append(row_data)
        
        columns = [
            'SI No', 'Date', 'Type', 'Invoice No', 'Store', 'Customer', 'Created By',
            'Product', 'Product Code', 'Quantity', 'Rate', 'Tax Type',
            'Tax 1 %', 'Tax 2 %', 'Tax Amount', 'Discount Amount', 'Amount', 'Final Amount'
        ]
        
        if is_return:
            columns.insert(4, 'Original Invoice')
        
        df = pd.DataFrame(data, columns=columns)
        return df

    def _export_to_format(self, df, file_format, title):
        try:
            output = BytesIO()
            
            if file_format == 'PDF':
                from reportlab.lib.pagesizes import letter, landscape
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib import colors
                from reportlab.lib.units import inch
                from reportlab.platypus.flowables import KeepTogether
                
                PRIMARY_COLOR = colors.HexColor("#250feb")
                WHITE = colors.white
                BLACK = colors.black
                GRID_COLOR = colors.HexColor('#e0e0e0')
                
                custom_size = (18 * inch, 16 * inch)
                
                doc = SimpleDocTemplate(
                    output,
                    pagesize=custom_size,
                    leftMargin=0.7*inch,
                    rightMargin=0.7*inch,
                    topMargin=0.7*inch,
                    bottomMargin=0.6*inch
                )
                elements = []
                
                styles = getSampleStyleSheet()
                title_style = styles['Title']
                title_style.textColor = PRIMARY_COLOR
                title_style.fontSize = 16
                title_style.alignment = 1
                
                elements.append(Paragraph(f"<b>{title}</b>", title_style))
                elements.append(Spacer(1, 12))
                
                if df.empty:
                    elements.append(Paragraph("No data available for the selected period", styles['Normal']))
                else:
                    formatted_data = []
                    for _, row in df.iterrows():
                        formatted_row = [str(row[col]) for col in df.columns]
                        formatted_data.append(formatted_row)
                    
                    data = [list(df.columns)] + formatted_data
                    
                    if 'Final Amount' in df.columns:
                        total_amount = df['Final Amount'].sum()
                        totals = [''] * len(df.columns)
                        totals[0] = "Total"
                        totals[-1] = f"{total_amount:,.2f}"
                        data.append(totals)
                    elif 'Amount' in df.columns:
                        total_amount = df['Amount'].sum()
                        totals = [''] * len(df.columns)
                        totals[0] = "Total"
                        totals[-1] = f"{total_amount:,.2f}"
                        data.append(totals)
                    
                    col_widths = [0.5*inch]  # SI No
                    col_widths += [0.8*inch]  # Date
                    col_widths += [1.3*inch]  # Type
                    col_widths += [1.2*inch]  # Invoice No
                    
                    if 'Original Invoice' in df.columns:
                        col_widths += [1.2*inch]  # Original Invoice
                    
                    col_widths += [
                        0.8*inch,  # Store
                        1.0*inch,  # Customer
                        1.0*inch,  # Created By
                        1.2*inch,  # Product
                        0.9*inch,  # Product Code
                        0.7*inch,  # Quantity
                        0.9*inch,  # Rate
                        0.9*inch,  # Tax Type
                        0.7*inch,  # Tax 1 %
                        0.7*inch,  # Tax 2 %
                        0.9*inch,  # Tax Amount
                        1.2*inch,  # Discount Amount
                        0.9*inch,  # Amount
                        1.0*inch,  # Final Amount
                    ]
                    
                    table = Table(data, colWidths=col_widths, repeatRows=1)
                    
                    style = TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), PRIMARY_COLOR),
                        ('TEXTCOLOR', (0, 0), (-1, 0), WHITE),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                        ('TOPPADDING', (0, 0), (-1, 0), 6),
                        ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # SI No
                        ('ALIGN', (1, 1), (1, -1), 'CENTER'),  # Date
                        ('ALIGN', (2, 1), (2, -1), 'CENTER'),  # Type
                        ('ALIGN', (3, 1), (4 if 'Original Invoice' in df.columns else 3, -1), 'LEFT'),  # Invoice No and Original Invoice
                        ('ALIGN', (5 if 'Original Invoice' in df.columns else 4, 1), (7 if 'Original Invoice' in df.columns else 6, -1), 'LEFT'),  # Store to Created By
                        ('ALIGN', (8 if 'Original Invoice' in df.columns else 7, 1), (-2, -1), 'RIGHT'),  # Numeric columns
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 9),
                        ('GRID', (0, 0), (-1, -1), 0.5, GRID_COLOR),
                        ('LINEBELOW', (0, 0), (-1, 0), 1, WHITE),
                        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                        ('LINEABOVE', (0, -1), (-1, -1), 1, BLACK),
                    ])
                    
                    table.setStyle(style)
                    elements.append(KeepTogether(table))
                
                doc.build(elements)
            
            elif file_format == 'DOCX':
                from docx import Document
                from docx.shared import Pt, RGBColor, Inches
                from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                
                document = Document()
                
                title_para = document.add_heading(title, level=1)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_para.runs[0]
                title_run.font.color.rgb = RGBColor(74, 111, 220)
                title_run.font.size = Pt(14)
                
                if df.empty:
                    document.add_paragraph("No data available for the selected period")
                else:
                    col_count = len(df.columns)
                    table = document.add_table(rows=1, cols=col_count)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Calculate column widths based on content
                    widths = []
                    for col in df.columns:
                        max_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
                        widths.append(min(max_len * 0.15, 2.5))  # Cap at 2.5 inches
                    
                    # Apply calculated widths
                    for i, width in enumerate(widths[:col_count]):
                        for cell in table.columns[i].cells:
                            cell.width = Inches(width)
                    
                    # Set header row
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(df.columns[:col_count]):
                        hdr_cells[i].text = str(header)
                        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        hdr_run = hdr_cells[i].paragraphs[0].runs[0]
                        hdr_run.font.bold = True
                        hdr_run.font.color.rgb = RGBColor(255, 255, 255)
                        hdr_run.font.size = Pt(10)
                        
                        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), '4a6fdc')
                        tcPr.append(shading)
                    
                    # Add data rows
                    for _, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for i, col in enumerate(df.columns[:col_count]):
                            value = row[col]
                            if isinstance(value, float):
                                value = f"{value:.2f}"
                            row_cells[i].text = str(value)
                            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            
                            if i in [0, 1, 2]:  # SI No, Date, Type
                                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif i in [3, 4] if 'Original Invoice' in df.columns else [3]:  # Invoice No and Original Invoice
                                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            elif isinstance(row[col], (int, float)) or col in ['Quantity', 'Rate', 'Tax 1 %', 
                                                                             'Tax 2 %', 'Tax Amount', 'Discount Amount', 
                                                                             'Amount', 'Final Amount']:
                                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                            for paragraph in row_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                    
                    # Add totals row if applicable
                    if 'Final Amount' in df.columns or 'Amount' in df.columns:
                        totals_row = table.add_row().cells
                        total_col = df.columns.get_loc('Final Amount') if 'Final Amount' in df.columns else df.columns.get_loc('Amount')
                        
                        for i in range(col_count):
                            if i >= col_count:
                                break
                            totals_row[i].text = ""
                            if i == 0:
                                totals_row[i].text = "Total"
                            elif i == total_col:
                                total_amount = df.iloc[:, total_col].sum()
                                totals_row[i].text = f"{total_amount:,.2f}"
                        
                        for cell in totals_row:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                
                document.save(output)
            
            elif file_format == 'XLSX':
                from xlsxwriter.utility import xl_rowcol_to_cell, xl_range
                
                with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                    if df.empty:
                        empty_df = pd.DataFrame({'Message': ['No data available for the selected period']})
                        empty_df.to_excel(writer, sheet_name=title[:31], index=False)
                    else:
                        # Create a copy to avoid modifying the original dataframe
                        report_df = df.copy()
                        
                        # Define numeric columns based on report type
                        numeric_cols = ['Quantity', 'Rate', 'Tax 1 %', 'Tax 2 %', 
                                      'Tax Amount', 'Discount Amount', 'Amount', 'Final Amount']
                        
                        # Clean and convert numeric columns
                        for col in numeric_cols:
                            if col in report_df.columns:
                                report_df[col] = pd.to_numeric(report_df[col], errors='coerce')
                                report_df[col] = report_df[col].fillna(0).apply(
                                    lambda x: f"{float(x):.2f}" if pd.notna(x) else "0.00"
                                )
                        
                        # Write data to Excel
                        report_df.to_excel(writer, sheet_name=title[:31], index=False)
                        
                        workbook = writer.book
                        worksheet = writer.sheets[title[:31]]
                        
                        # Define formats
                        header_format = workbook.add_format({
                            'bold': True,
                            'font_color': 'white',
                            'bg_color': '#4a6fdc',
                            'align': 'center',
                            'valign': 'vcenter',
                            'border': 1,
                            'font_size': 10
                        })
                        
                        text_format = workbook.add_format({
                            'align': 'left',
                            'valign': 'vcenter',
                            'border': 1,
                            'font_size': 9
                        })
                        
                        num_format = workbook.add_format({
                            'align': 'right',
                            'valign': 'vcenter',
                            'border': 1,
                            'num_format': '#,##0.00',
                            'font_size': 9
                        })
                        
                        center_format = workbook.add_format({
                            'align': 'center',
                            'valign': 'vcenter',
                            'border': 1,
                            'font_size': 9
                        })
                        
                        total_format = workbook.add_format({
                            'bold': True,
                            'align': 'right',
                            'valign': 'vcenter',
                            'border': 1,
                            'num_format': '#,##0.00',
                            'font_size': 9
                        })
                        
                        summary_format = workbook.add_format({
                            'bold': True,
                            'align': 'left',
                            'valign': 'vcenter',
                            'border': 1,
                            'font_size': 9
                        })
                        
                        # Apply header format
                        worksheet.set_row(0, 20, header_format)
                        
                        # Write data with proper formatting
                        for idx, row in enumerate(report_df.itertuples(), 1):
                            worksheet.set_row(idx, 18)
                            for col_idx, col in enumerate(report_df.columns, 0):
                                value = row[col_idx + 1]  # itertuples index starts at 1
                                
                                if col_idx == 0:  # SI No
                                    worksheet.write(idx, col_idx, value, center_format)
                                elif col == 'Type':  # Type column
                                    worksheet.write(idx, col_idx, value, center_format)
                                elif col in numeric_cols:  # Numeric columns
                                    try:
                                        num_value = float(value) if value and str(value).strip() else 0
                                        worksheet.write_number(idx, col_idx, num_value, num_format)
                                    except (ValueError, TypeError):
                                        worksheet.write(idx, col_idx, value, text_format)
                                else:  # Text columns
                                    worksheet.write(idx, col_idx, value, text_format)
                        
                        # Add totals row if applicable
                        if 'Final Amount' in report_df.columns:
                            amount_col = report_df.columns.get_loc('Final Amount')
                            worksheet.write_formula(
                                len(report_df)+1, 0,
                                "Total", summary_format
                            )
                            worksheet.write_formula(
                                len(report_df)+1, amount_col,
                                f'=SUM({xl_rowcol_to_cell(1, amount_col)}:{xl_rowcol_to_cell(len(report_df), amount_col)})',
                                total_format
                            )
                        elif 'Amount' in report_df.columns:
                            amount_col = report_df.columns.get_loc('Amount')
                            worksheet.write_formula(
                                len(report_df)+1, 0,
                                "Total", summary_format
                            )
                            worksheet.write_formula(
                                len(report_df)+1, amount_col,
                                f'=SUM({xl_rowcol_to_cell(1, amount_col)}:{xl_rowcol_to_cell(len(report_df), amount_col)})',
                                total_format
                            )
                        
                        # Auto-adjust column widths
                        for i, col in enumerate(report_df.columns):
                            max_len = max(
                                report_df[col].astype(str).map(len).max(),
                                len(str(col))
                            ) + 2
                            worksheet.set_column(i, i, max_len)
            
            output.seek(0)
            return output
            
        except Exception as e:
            logger.error(f"Error generating report file: {str(e)}", exc_info=True)
            raise serializers.ValidationError(f"Error generating report file: {str(e)}")